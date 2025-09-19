#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Remote-sensing caption batcher with GLM/Qwen (OpenAI-compatible) backends.
- Recursively scans a folder for images/arrays
- Optional HSI panelization (band triplets -> RGB panels)
- Sends single/multi images in one user message
- Robust JSON extraction and retry
- Checkpoint/resume via JSONL
- Multi-threaded concurrent API requests for speed
All comments are in English as requested.
"""

import argparse
import base64
import json
import os
import re
import time
import tempfile
import threading
from itertools import islice
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed
from queue import Queue

import numpy as np
import requests
import cv2
from PIL import Image, UnidentifiedImageError
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document

# =========================
# Document processing utilities
# =========================
def extract_pdf_content(pdf_path: Path, max_pages: int = 50) -> str:
    """提取PDF文本内容，限制页数避免过长"""
    try:
        reader = PdfReader(pdf_path)
        content = ""
        total_pages = min(len(reader.pages), max_pages)
        for i in range(total_pages):
            page_text = reader.pages[i].extract_text()
            content += f"\n--- 第{i+1}页 ---\n{page_text}"
        if len(reader.pages) > max_pages:
            content += f"\n\n[注: 文档共{len(reader.pages)}页，仅显示前{max_pages}页]"
        return content
    except Exception as e:
        return f"PDF读取错误: {str(e)}"

def extract_csv_content(csv_path: Path, max_rows: int = 200) -> str:
    """提取CSV内容，包含统计信息"""
    try:
        df = pd.read_csv(csv_path)
        content = f"CSV文件统计信息:\n"
        content += f"- 行数: {len(df)}\n"
        content += f"- 列数: {len(df.columns)}\n"
        content += f"- 列名: {', '.join(df.columns.tolist())}\n\n"
        
        if len(df) > max_rows:
            content += f"数据预览 (前{max_rows}行):\n"
            content += df.head(max_rows).to_string(index=False)
            content += f"\n\n数据统计摘要:\n{df.describe().to_string()}"
        else:
            content += f"完整数据:\n{df.to_string(index=False)}"
        
        return content
    except Exception as e:
        return f"CSV读取错误: {str(e)}"

def extract_txt_content(txt_path: Path, max_chars: int = 10000) -> str:
    """提取TXT文本内容"""
    try:
        with open(txt_path, 'r', encoding='utf-8') as f:
            content = f.read()
        if len(content) > max_chars:
            content = content[:max_chars] + f"\n\n[注: 文档已截断，原长度{len(content)}字符]"
        return content
    except Exception as e:
        return f"TXT读取错误: {str(e)}"

def extract_docx_content(docx_path: Path, max_paragraphs: int = 100) -> str:
    """提取DOCX文档内容"""
    try:
        doc = Document(docx_path)
        content = ""
        for i, para in enumerate(doc.paragraphs):
            if i >= max_paragraphs:
                content += f"\n\n[注: 文档共{len(doc.paragraphs)}段，仅显示前{max_paragraphs}段]"
                break
            content += para.text + "\n"
        return content
    except Exception as e:
        return f"DOCX读取错误: {str(e)}"

def process_reference_documents(ref_docs: List[Path]) -> str:
    """处理所有参考文档并合并内容"""
    if not ref_docs:
        return ""
    
    combined_content = "=== 参考文档内容 ===\n\n"
    
    for i, doc_path in enumerate(ref_docs, 1):
        if not doc_path.exists():
            combined_content += f"{i}. {doc_path.name}: 文件不存在\n\n"
            continue
            
        ext = doc_path.suffix.lower()
        combined_content += f"{i}. 文档: {doc_path.name}\n"
        combined_content += "-" * 50 + "\n"
        
        if ext == '.pdf':
            content = extract_pdf_content(doc_path)
        elif ext == '.csv':
            content = extract_csv_content(doc_path)
        elif ext in ['.txt', '.md']:
            content = extract_txt_content(doc_path)
        elif ext == '.docx':
            content = extract_docx_content(doc_path)
        else:
            content = f"不支持的文档类型: {ext}"
        
        combined_content += content + "\n\n"
    
    combined_content += "=== 参考文档结束 ===\n\n"
    return combined_content



# =========================
# Defaults
# =========================
DEFAULT_API_URL = "http://localhost:8000/v1/chat/completions"
DEFAULT_MODEL = "glm-4.5v"
DEFAULT_EXTS = (".jpg", ".jpeg", ".png", ".tif", ".tiff", ".npy", ".npz")
DEFAULT_MAX_SIDE = 2048
DEFAULT_TEMPERATURE = 0.2
DEFAULT_TOP_P = 0.9
DEFAULT_MAX_TOKENS = 1500
DEFAULT_LIMIT = 0
DEFAULT_MAX_BYTES = 0
DEFAULT_MAX_PIXELS = 0
DEFAULT_API_TIMEOUT = 180
DEFAULT_MAX_RETRIES = 5
DEFAULT_MAX_WORKERS = 4  # New: default concurrent workers
HEADERS = {"Content-Type": "application/json"}

# Thread-safe write lock
write_lock = threading.Lock()

# =========================
# Prompt helpers
# =========================
def load_prompts(json_path: Path, prompt_name: str) -> Tuple[str, str]:
    """Load {system,user} from a JSON file by name; user may be a list joined to str."""
    if not json_path.exists():
        raise FileNotFoundError(f"Prompt file not found: {json_path}")
    data = json.loads(json_path.read_text(encoding="utf-8"))
    if prompt_name not in data:
        raise KeyError(f"Prompt '{prompt_name}' not in {json_path}")
    sys_prompt = data[prompt_name].get("system", "")
    user_prompt = data[prompt_name].get("user", "")
    if isinstance(user_prompt, list):
        user_prompt = "".join(user_prompt)
    if not sys_prompt or not user_prompt:
        raise ValueError(f"Prompt '{prompt_name}' missing system or user")
    return sys_prompt, user_prompt

# =========================
# File & image utilities
# =========================
def iter_images_stream(root: Path, exts: Tuple[str, ...]) -> Iterable[Path]:
    """Yield files recursively matching extensions."""
    exts_set = set(x.lower() for x in exts)
    for dirpath, _, files in os.walk(root):
        for fn in files:
            if any(fn.lower().endswith(e) for e in exts_set):
                yield Path(dirpath) / fn

def file_too_large(p: Path, max_bytes: int) -> bool:
    """Size filter in bytes."""
    if max_bytes > 0:
        try:
            return p.stat().st_size > max_bytes
        except OSError:
            return False
    return False

def image_too_many_pixels(p: Path, max_pixels: int) -> bool:
    """Pixel-count filter via PIL; fail-safe to False."""
    if max_pixels > 0:
        try:
            with Image.open(p) as img:
                w, h = img.size
            return (w * h) > max_pixels
        except Exception:
            return False
    return False

def load_image_size(path: Path) -> Tuple[int, int]:
    """Return (w, h) with PIL."""
    with Image.open(path) as img:
        return img.size

def save_resized_if_needed(path: Path, max_side: int, outdir: Path) -> Tuple[Path, Tuple[int, int], Tuple[int, int]]:
    """Resize to fit max_side while keeping aspect ratio; write only if needed."""
    ow, oh = load_image_size(path)
    if max(ow, oh) <= max_side:
        return path, (ow, oh), (ow, oh)
    out = outdir / f"{path.stem}_resized{path.suffix}"
    with Image.open(path) as img:
        img = img.convert("RGB") if img.mode not in ("RGB", "L") else img
        img.thumbnail((max_side, max_side))
        img.save(out)
    nw, nh = load_image_size(out)
    return out, (ow, oh), (nw, nh)

def load_processed_paths(jsonl_path: Path) -> set:
    """Recover processed image paths from a possibly pretty-printed JSONL (brace-balanced)."""
    done = set()
    if not jsonl_path.exists():
        return done
    buf, depth = "", 0
    with jsonl_path.open("r", encoding="utf-8") as f:
        for line in f:
            buf += line
            depth += line.count("{") - line.count("}")
            if depth == 0 and buf.strip():
                try:
                    obj = json.loads(buf)
                    ip = obj.get("image_path")
                    if ip:
                        done.add(str(Path(ip).resolve()))
                except Exception:
                    pass
                buf = ""
    if depth == 0 and buf.strip():
        try:
            obj = json.loads(buf)
            ip = obj.get("image_path")
            if ip:
                done.add(str(Path(ip).resolve()))
        except Exception:
            pass
    return done

# =========================
# JSON extraction
# =========================
def try_json_load(s: str) -> Optional[Any]:
    try:
        return json.loads(s)
    except Exception:
        return None

def balanced_json_extract(text: str) -> Optional[Any]:
    """Extract a top-level balanced {...} block from free text."""
    s = text.strip()
    if "{" not in s:
        return None
    start = s.find("{")
    depth = 0
    for i in range(start, len(s)):
        ch = s[i]
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                cand = s[start : i + 1]
                cand = re.sub(r",\s*([}\]])", r"\1", cand)
                obj = try_json_load(cand)
                if obj is not None:
                    return obj
    return None

def robust_parse_json(text: str) -> Optional[Any]:
    """Try direct -> fenced ```json -> balanced block."""
    obj = try_json_load(text.strip())
    if obj is not None:
        return obj
    m = re.search(r"```json\s*(\{.*?\})\s*```", text, flags=re.S | re.M)
    if m:
        obj = try_json_load(m.group(1))
        if obj is not None:
            return obj
    return balanced_json_extract(text)

# =========================
# Hyperspectral helpers
# =========================
def _safe_import_tifffile():
    try:
        import tifffile  # type: ignore
        return tifffile
    except Exception:
        return None

def _safe_import_rasterio():
    try:
        import rasterio  # type: ignore
        return rasterio
    except Exception:
        return None

def _percentile_stretch(arr: np.ndarray, clip_percent: float = 2.0) -> np.ndarray:
    """Per-channel percentile stretch to uint8."""
    arr = arr.astype(np.float32)
    if arr.ndim == 2:
        arr = arr[..., None]
    out = np.zeros_like(arr, dtype=np.uint8)
    for c in range(arr.shape[-1]):
        ch = arr[..., c]
        lo = np.percentile(ch, clip_percent)
        hi = np.percentile(ch, 100 - clip_percent)
        if hi <= lo:
            lo, hi = ch.min(), ch.max()
            if hi <= lo:
                out[..., c] = 0
                continue
        ch = np.clip((ch - lo) / (hi - lo + 1e-8), 0, 1) * 255.0
        out[..., c] = ch.astype(np.uint8)
    return out[..., 0] if out.shape[-1] == 1 else out

def _load_array_from_path(p: Path) -> np.ndarray:
    """Load npy/npz, GeoTIFF (rasterio/tifffile), or common images; (C,H,W)->(H,W,C)."""
    suf = p.suffix.lower()
    if suf == ".npy":
        arr = np.load(p)
        if arr.ndim == 3 and arr.shape[0] < arr.shape[-1]:
            arr = np.transpose(arr, (1, 2, 0))
        return arr
    if suf == ".npz":
        data = np.load(p)
        key = list(data.keys())[0]
        arr = data[key]
        if arr.ndim == 3 and arr.shape[0] < arr.shape[-1]:
            arr = np.transpose(arr, (1, 2, 0))
        return arr
    rio = _safe_import_rasterio()
    if rio is not None and suf in (".tif", ".tiff"):
        try:
            with rio.open(p.as_posix()) as ds:
                arr = ds.read()  # (C,H,W)
                return np.transpose(arr, (1, 2, 0))
        except Exception:
            pass
    tiff = _safe_import_tifffile()
    if tiff is not None and suf in (".tif", ".tiff"):
        try:
            arr = tiff.imread(p.as_posix())
            if arr.ndim == 3 and arr.shape[0] < arr.shape[-1] and arr.shape[0] <= 64:
                arr = np.transpose(arr, (1, 2, 0))
            return arr
        except Exception:
            pass
    # Fallback to PIL/CV2
    try:
        with Image.open(p) as img:
            return np.array(img)
    except Exception:
        img = cv2.imread(p.as_posix(), cv2.IMREAD_UNCHANGED)
        if img is None:
            raise ValueError(f"Cannot read file: {p}")
        if img.ndim == 3 and img.shape[2] == 3:
            img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
        return img

def _parse_band_indices(spec: str, indexing: str, total_bands: int) -> Tuple[int, int, int]:
    """Parse 'a,b,c' to 0-based indices per chosen indexing."""
    parts = [s.strip() for s in spec.split(",") if s.strip()]
    if len(parts) != 3:
        raise ValueError(f"Triplet must be 'a,b,c': {spec}")
    idx = [int(x) for x in parts]
    if indexing == "one-based":
        idx = [i - 1 for i in idx]
    for i in idx:
        if i < 0 or i >= total_bands:
            raise ValueError(f"Band {i} out of range [0,{total_bands-1}]")
    return tuple(idx)  # type: ignore

def _mk_panels_from_hsi(arr: np.ndarray,
                        triplets: str,
                        indexing: str,
                        clip_percent: float,
                        max_panels: int) -> List[Tuple[str, np.ndarray]]:
    """Make multiple (label, RGB_u8) panels from an HSI cube using triplets spec."""
    if arr.ndim == 2:
        arr = arr[..., None]
    H, W, C = arr.shape
    panels: List[Tuple[str, np.ndarray]] = []
    for gi, gspec in enumerate([g for g in triplets.split(";") if g.strip()], 1):
        if len(panels) >= max_panels:
            break
        try:
            r, g, b = _parse_band_indices(gspec, indexing, C)
            cube3 = np.stack([arr[..., r], arr[..., g], arr[..., b]], axis=-1)
            rgb = _percentile_stretch(cube3, clip_percent)
            panels.append((f"P{gi}({gspec})", rgb))
        except Exception:
            continue
    return panels

def _write_temp_png(rgb_u8: np.ndarray, outdir: Path, stem: str, suffix: str) -> Path:
    """Write a uint8 RGB array to PNG."""
    safe = suffix.replace("(", "_").replace(")", "").replace(",", "-").replace(" ", "_")
    outp = outdir / f"{stem}_{safe}.png"
    Image.fromarray(rgb_u8).save(outp)
    return outp

# =========================
# Backend-agnostic caller
# =========================
def _encode_b64(p: Path) -> str:
    with open(p, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")

def _image_part_from_path(p: Path, transport: str) -> Dict[str, Any]:
    """Return an OpenAI-compatible image content part."""
    if transport == "dataurl":
        return {"type": "image_url", "image_url": {"url": f"data:image/*;base64,{_encode_b64(p)}"}}
    return {"type": "image_url", "image_url": {"url": f"file://{p.resolve()}"}}

def call_openai_compatible(
    api_url: str,
    model: str,
    images: List[Path],
    system_prompt: str,
    user_prompt: str,
    temperature: float,
    top_p: float,
    max_tokens: int,
    json_mode: bool,
    timeout: int = DEFAULT_API_TIMEOUT,
    transport: str = "file",
    reference_content: str = "",
) -> Dict[str, Any]:
    """Single entry for GLM/Qwen (OpenAI-compatible) backends."""
    full_user_prompt = user_prompt
    if reference_content.strip():
        full_user_prompt = reference_content + "\n" + user_prompt

    content = [{"type": "text", "text": full_user_prompt}] + [_image_part_from_path(p, transport) for p in images]
    payload: Dict[str, Any] = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": content},
        ],
        "temperature": temperature,
        "top_p": top_p,
        "max_tokens": max_tokens,
    }
    if json_mode:
        payload["response_format"] = {"type": "json_object"}
    t0 = time.time()
    resp = requests.post(api_url, headers=HEADERS, data=json.dumps(payload), timeout=timeout)
    dt = time.time() - t0
    if resp.status_code != 200:
        return {"status": resp.status_code, "elapsed_sec": dt, "error_text": resp.text[:2000]}
    data = resp.json()
    try:
        content = data["choices"][0]["message"]["content"]
    except Exception:
        return {"status": 200, "elapsed_sec": dt, "error_text": str(data)[:2000]}
    parsed = robust_parse_json(content)
    return {"elapsed_sec": dt, "raw_text": content, "json": parsed}

# =========================
# Task definition for worker threads
# =========================
class ProcessingTask:
    def __init__(self, idx: int, img_path: Path, file_subdir: Path, 
                 panels_info: List[Tuple[str, Path]], send_paths: List[Path], 
                 use_single_path: Optional[Path], ow: int, oh: int, iw: int, ih: int, 
                 mae_value: float, user_prompt: str, reference_content: str = ""):
        self.idx = idx
        self.img_path = img_path
        self.file_subdir = file_subdir
        self.panels_info = panels_info
        self.send_paths = send_paths
        self.use_single_path = use_single_path
        self.ow = ow
        self.oh = oh
        self.iw = iw
        self.ih = ih
        self.mae_value = mae_value
        self.user_prompt = user_prompt
        self.reference_content = reference_content 

# =========================
# Worker function for concurrent processing
# =========================
def process_single_image(
    task: ProcessingTask,
    api_url: str,
    model: str,
    system_prompt: str,
    temperature: float,
    top_p: float,
    max_tokens: int,
    json_mode: bool,
    transport: str,
    max_retries: int,
    verbose: bool,
    detail: bool,
    total_files: int,
) -> Dict[str, Any]:
    """Process a single image task and return the result record."""
    
    # Retry wrapper
    retry = 0
    result = None
    while retry <= max_retries:
        send = task.send_paths if task.panels_info else [task.use_single_path]  # type: ignore
        out = call_openai_compatible(
            api_url=api_url,
            model=model,
            images=send,
            system_prompt=system_prompt,
            user_prompt=task.user_prompt,
            temperature=temperature,
            top_p=top_p,
            max_tokens=max_tokens,
            json_mode=json_mode,
            transport=transport,
            reference_content=task.reference_content,
        )
        if out.get("json") is not None:
            result = out
            break
        retry += 1
        if retry <= max_retries:
            print(f"[RETRY {retry}/{max_retries}] Null JSON, retrying {task.img_path.name}...")
        else:
            result = out
            print(f"[FAILED] Max retries reached for {task.img_path.name}")

    if verbose and result.get("raw_text"):
        print(f"[RAW] {'-'*40}\n{result['raw_text']}\n{'-'*40}")

    record: Dict[str, Any] = {
        "image_path": str(task.img_path),
        "orig_size": {"w": task.ow, "h": task.oh},
        "infer_size": {"w": task.iw, "h": task.ih},
        "high_freq_restore_diff": task.mae_value,
        "elapsed_sec": result.get("elapsed_sec", None),
        "retry_count": retry,
        "panels": [{"label": lbl, "path": str(p)} for (lbl, p) in task.panels_info] if task.panels_info else None,
        "output": {
            "global_caption": result.get("json", {}).get("global_caption", ""),
            "event_analysis": result.get("json", {}).get("event_analysis", ""),
            "spatial_attributes": result.get("json", {}).get("spatial_attributes", ""),
            "spectral_features": result.get("json", {}).get("spectral_features", ""),
            "task_relevance": result.get("json", {}).get("task_relevance", "")
        } if result.get("json") else None
    }

    if detail and result.get("raw_text"):
        record["raw_text"] = result["raw_text"]
    if "error_text" in result:
        record["error"] = result.get("error_text", None)
    
    # Progress reporting
    progress = (task.idx / total_files * 100) if total_files > 0 else 0
    required_keys = ["global_caption", "event_analysis", "spatial_attributes", "spectral_features", "task_relevance"]
    
    # start time to time
    elapsed_time = time.time() - getattr(process_single_image, 'start_time', time.time())
    
    # format progress
    print(f"[{task.idx:06d}/{progress:.1f}%/{elapsed_time:.1f}s] {task.img_path.name}\t{record['elapsed_sec']:.2f}s", end="\t")
    if record["output"] is None:
        print("[MISS] All output keys")
    else:
        miss = [k for k in required_keys if not record["output"].get(k)]
        print("[OK] All output keys present" if not miss else f"[MISS] {', '.join(miss)}")

    return record

# =========================
# Processing loop with concurrency
# =========================
def _process_stream_concurrent(
    stream: Iterable[Path],
    total_files: int,
    tmp_root: Path,
    reference_docs: List[Path] = None, 
    *,
    api_url: str,
    model: str,
    max_side: int,
    out_jsonl: Path,
    temperature: float,
    top_p: float,
    max_tokens: int,
    max_bytes: int,
    max_pixels: int,
    detail: bool,
    system_prompt: str,
    user_prompt_template: str,
    verbose: bool,
    max_retries: int,
    hsi_panels: bool,
    hsi_triplets: str,
    band_indexing: str,
    clip_percent: float,
    max_panels: int,
    processed: set,
    json_mode: bool,
    transport: str,
    max_workers: int,
) -> Tuple[int, int]:
    """Core concurrent processing loop."""

    reference_content = ""
    if reference_docs:
        reference_content = process_reference_documents(reference_docs)
        print(f"[INFO] Loaded {len(reference_docs)} reference document(s), total length: {len(reference_content)} chars")

    total_seen = 0
    total_emitted = 0
    
    # Process tasks concurrently with streaming approach
    with open(out_jsonl, "a", encoding="utf-8") as fout:
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Use a queue to limit the number of pending futures
            future_queue = set()
            
            for idx, img_path in enumerate(stream, 1):
                if str(img_path.resolve()) in processed:
                    print(f"[SKIP] Already processed: {img_path.name}")
                    continue
                total_seen += 1

                if file_too_large(img_path, max_bytes):
                    print(f"[SKIP] Too large (>max-bytes): {img_path.name}")
                    continue
                if image_too_many_pixels(img_path, max_pixels):
                    print(f"[SKIP] Too many pixels (>max-pixels): {img_path.name}")
                    continue

                try:
                    panels_info: List[Tuple[str, Path]] = []
                    send_paths: List[Path] = []
                    ow = oh = iw = ih = None  # type: Optional[int]
                    mae_value = -1.0
                    use_single_path: Optional[Path] = None

                    # Try loading to detect multichannel HSI
                    try:
                        arr = _load_array_from_path(img_path)
                    except Exception:
                        arr = None
                    is_multi_channel = (arr is not None and arr.ndim == 3 and arr.shape[2] > 3)

                    # Create a per-file subdir (keeps artifacts tidy)
                    file_subdir = tmp_root / img_path.stem
                    file_subdir.mkdir(parents=True, exist_ok=True)

                    if hsi_panels and is_multi_channel:
                        panels = _mk_panels_from_hsi(arr, hsi_triplets, band_indexing, clip_percent, max_panels)
                        if len(panels) == 0:
                            use_path, (ow, oh), (iw, ih) = save_resized_if_needed(img_path, max_side, file_subdir)
                            use_single_path = use_path
                        else:
                            H, W, _ = arr.shape
                            ow, oh = W, H
                            for label, rgb in panels:
                                png = _write_temp_png(rgb, file_subdir, img_path.stem, label)
                                use_path, (_ow2, _oh2), (riw, rih) = save_resized_if_needed(png, max_side, file_subdir)
                                send_paths.append(use_path)
                                panels_info.append((label, use_path))
                            if send_paths:
                                try:
                                    # Use first panel for the fast MAE proxy
                                    gray_mae_src = send_paths[0]
                                    img_gray = cv2.imread(str(gray_mae_src), cv2.IMREAD_GRAYSCALE)
                                    dft = cv2.dft(np.float32(img_gray), flags=cv2.DFT_COMPLEX_OUTPUT)
                                    rows, cols = img_gray.shape
                                    crow, ccol = rows // 2, cols // 2
                                    mask = np.zeros((rows, cols, 2), np.uint8)
                                    radius = int(min(rows, cols) * 0.5)
                                    cv2.circle(mask, (ccol, crow), radius, (1, 1), -1)
                                    filtered = dft * mask
                                    back = cv2.idft(filtered, flags=cv2.DFT_SCALE | cv2.DFT_REAL_OUTPUT)
                                    mae_value = float(np.mean(np.abs(img_gray - back)))
                                except Exception:
                                    mae_value = -1.0
                                iw, ih = load_image_size(send_paths[0])
                    else:
                        use_path, (ow, oh), (iw, ih) = save_resized_if_needed(img_path, max_side, file_subdir)
                        use_single_path = use_path
                        try:
                            img_gray = cv2.imread(str(use_single_path), cv2.IMREAD_GRAYSCALE)
                            dft = cv2.dft(np.float32(img_gray), flags=cv2.DFT_COMPLEX_OUTPUT)
                            rows, cols = img_gray.shape
                            crow, ccol = rows // 2, cols // 2
                            mask = np.zeros((rows, cols, 2), np.uint8)
                            radius = int(min(rows, cols) * 0.5)
                            cv2.circle(mask, (ccol, crow), radius, (1, 1), -1)
                            filtered = dft * mask
                            back = cv2.idft(filtered, flags=cv2.DFT_SCALE | cv2.DFT_REAL_OUTPUT)
                            mae_value = float(np.mean(np.abs(img_gray - back)))
                        except Exception:
                            mae_value = -1.0

                    if None in (ow, oh, iw, ih):
                        if arr is not None and arr.ndim >= 2:
                            _w, _h = arr.shape[1], arr.shape[0]
                        else:
                            _w, _h = load_image_size(img_path)
                        ow, oh, iw, ih = _w, _h, _w, _h

                    panel_lines = ""
                    if panels_info:
                        panel_lines = "Panels:\n" + "\n".join([f"- {i+1}. {lbl} -> {p.name}" for i, (lbl, p) in enumerate(panels_info)]) + "\n"

                    user_prompt = (
                        f"{panel_lines}"
                        f"Image original size: {ow}x{oh} pixels, resized to: {iw}x{ih} pixels\n"
                        f"High frequency restoration difficulty: {mae_value:.2f} (use this to adapt detail level)\n"
                        + user_prompt_template.format(orig_w=ow, orig_h=oh, inf_w=iw, inf_h=ih)
                    )

                    # Create task
                    task = ProcessingTask(
                        idx, img_path, file_subdir, panels_info, send_paths, 
                        use_single_path, ow, oh, iw, ih, mae_value, user_prompt,
                        reference_content  
                    )

                    # Submit task to executor immediately
                    future = executor.submit(
                        process_single_image,
                        task, api_url, model, system_prompt, temperature, top_p,
                        max_tokens, json_mode, transport, max_retries, verbose,
                        detail, total_files
                    )
                    future_queue.add(future)

                    # Remove completed futures and process results
                    completed = {f for f in future_queue if f.done()}
                    for future in completed:
                        future_queue.remove(future)
                        try:
                            record = future.result()
                            # Thread-safe write to file
                            with write_lock:
                                fout.write(json.dumps(record, ensure_ascii=False, indent=2) + "\n")
                                fout.flush()
                            total_emitted += 1
                        except Exception as e:
                            print(f"[ERROR] Processing {img_path.name}: {e}")
                            # Write error record
                            fallback = {"image_path": str(img_path), "error": f"exception: {repr(e)}"}
                            with write_lock:
                                fout.write(json.dumps(fallback, ensure_ascii=False) + "\n")
                                fout.flush()
                except UnidentifiedImageError:
                    print(f"[SKIP] Invalid/corrupted image: {img_path.name}")
                except Exception as e:
                    print(f"[ERROR] {img_path.name}: {e}")

            # Process any remaining futures
            for future in as_completed(future_queue):
                try:
                    record = future.result()
                    with write_lock:
                        fout.write(json.dumps(record, ensure_ascii=False, indent=2) + "\n")
                        fout.flush()
                    total_emitted += 1
                except Exception as e:
                    # Get task info from the exception if possible
                    print(f"[ERROR] Processing remaining task: {e}")
                    with write_lock:
                        fout.write(json.dumps({"error": f"exception: {repr(e)}"}, ensure_ascii=False) + "\n")
                        fout.flush()

    return total_seen, total_emitted

# =========================
# CLI
# =========================
def main():
    ap = argparse.ArgumentParser(description="Batch remote-sensing captioning with GLM/Qwen backends.")
    ap.add_argument("folder", type=str, help="Image folder (recursive)")
    ap.add_argument("--url", type=str, default=DEFAULT_API_URL)
    ap.add_argument("--model", type=str, default=DEFAULT_MODEL)
    ap.add_argument("--out", type=str, default="results.jsonl")
    ap.add_argument("--max-side", type=int, default=DEFAULT_MAX_SIDE)
    ap.add_argument("--exts", type=str, default=",".join(DEFAULT_EXTS))
    ap.add_argument("--temperature", type=float, default=DEFAULT_TEMPERATURE)
    ap.add_argument("--top-p", type=float, default=DEFAULT_TOP_P)
    ap.add_argument("--max-tokens", type=int, default=DEFAULT_MAX_TOKENS)
    ap.add_argument("--limit", type=int, default=DEFAULT_LIMIT)
    ap.add_argument("--max-bytes", type=int, default=DEFAULT_MAX_BYTES)
    ap.add_argument("--max-pixels", type=int, default=DEFAULT_MAX_PIXELS)
    ap.add_argument("--verbose", action="store_true")
    ap.add_argument("--prompts", type=str, default="prompts.json")
    ap.add_argument("--prompt-name", type=str, default="default")
    ap.add_argument("--detail", action="store_true")
    ap.add_argument("--overwrite", action="store_true")
    ap.add_argument("--max-retries", type=int, default=DEFAULT_MAX_RETRIES)
    
    ap.add_argument("--reference-docs", type=str, nargs="+", default=None,
                    help="参考文档路径列表，支持PDF、CSV、TXT、DOCX格式")
    ap.add_argument("--max-doc-pages", type=int, default=50,
                    help="PDF文档最大处理页数 (default: 50)")
    ap.add_argument("--max-csv-rows", type=int, default=200,
                    help="CSV文件最大显示行数 (default: 200)")
    ap.add_argument("--max-txt-chars", type=int, default=10000,
                    help="TXT文件最大字符数 (default: 10000)")

    # parallel
    ap.add_argument("--max-workers", type=int, default=DEFAULT_MAX_WORKERS,
                    help="Maximum number of concurrent API workers (default: 4)")

    # Backend & transport
    ap.add_argument("--backend", type=str, choices=["glm", "qwen"], default="glm",
                    help="Select backend for small compatibility tweaks (kept for clarity).")
    ap.add_argument("--image-transport", type=str, choices=["file", "dataurl"], default="file",
                    help="file: file:// path (needs shared FS); dataurl: base64 inline (safer for Qwen).")
    ap.add_argument("--json-mode", action="store_true",
                    help="Send response_format={'type':'json_object'} (recommended for Qwen).")

    # HSI options
    ap.add_argument("--hsi-panels", action="store_true",
                    help="Split HSI into multiple RGB panels and send together.")
    ap.add_argument("--hsi-triplets", type=str, default="4,3,2;8,4,3;12,8,4",
                    help="Semicolon-separated band triplets, e.g. '3,2,1;8,4,3'.")
    ap.add_argument("--band-indexing", type=str, default="one-based", choices=["one-based", "zero-based"])
    ap.add_argument("--clip-percent", type=float, default=2.0)
    ap.add_argument("--max-panels", type=int, default=6)
    ap.add_argument("--keep-panels-dir", type=str, default=None,
                    help="Persist panels/intermediates here (no auto-clean). Omit to use tempdir.")

    args = ap.parse_args()

    folder = Path(args.folder).expanduser().resolve()
    if not folder.is_dir():
        raise FileNotFoundError(f"Input folder not readable: {folder}")

    out_path = Path(args.out).expanduser().resolve()
    if args.overwrite and out_path.exists():
        print(f"[INFO] --overwrite: delete {out_path}")
        out_path.unlink()
    out_path.parent.mkdir(parents=True, exist_ok=True)

    exts = tuple(x.strip().lower() if x.strip().startswith(".") else "." + x.strip().lower()
                 for x in args.exts.split(",") if x.strip())
    prompts_path = Path(args.prompts).expanduser().resolve()
    system_prompt, user_prompt_template = load_prompts(prompts_path, args.prompt_name)
    
    # Count total files (with limit awareness)
    def count_files_fast(folder: Path, exts: Tuple[str, ...]) -> int:
        cnt, ex = 0, {e.lower() for e in exts}
        for dp, _, files in os.walk(folder):
            for f in files:
                if any(f.lower().endswith(s) for s in ex):
                    cnt += 1
        return cnt
    
    # add time record
    process_single_image.start_time = time.time()
    
    total_files = args.limit if args.limit and args.limit > 0 else count_files_fast(folder, exts)
    print(f"[INFO] Total files to process: {total_files}")

    # Resume set
    processed = load_processed_paths(out_path) if out_path.exists() else set()
    mode = "append" if processed else "create"
    print(f"[INFO] Results file mode: {mode}, already processed: {len(processed)}")
    print(f"[INFO] Using {args.max_workers} concurrent workers for API requests")

    # Build stream
    stream: Iterable[Path] = iter_images_stream(folder, exts)
    if args.limit and args.limit > 0:
        stream = islice(stream, args.limit)

    reference_docs = []
    if args.reference_docs:
        for doc_path in args.reference_docs:
            path = Path(doc_path).expanduser().resolve()
            if path.exists():
                reference_docs.append(path)
                print(f"[INFO] Added reference document: {path}")
            else:
                print(f"[WARN] Reference document not found: {path}")

    # Temp vs persistent artifacts dir
    if args.keep_panels_dir:
        tmp_root = Path(args.keep_panels_dir).expanduser().resolve()
        tmp_root.mkdir(parents=True, exist_ok=True)
        print(f"[INFO] Using persistent dir: {tmp_root}")
        seen, emitted = _process_stream_concurrent(
            stream, total_files, tmp_root, reference_docs,
            api_url=args.url, model=args.model, max_side=args.max_side, out_jsonl=out_path,
            temperature=args.temperature, top_p=args.top_p, max_tokens=args.max_tokens,
            max_bytes=args.max_bytes, max_pixels=args.max_pixels, detail=args.detail,
            system_prompt=system_prompt, user_prompt_template=user_prompt_template,
            verbose=args.verbose, max_retries=args.max_retries,
            hsi_panels=args.hsi_panels, hsi_triplets=args.hsi_triplets,
            band_indexing=args.band_indexing, clip_percent=args.clip_percent,
            max_panels=args.max_panels, processed=processed, json_mode=args.json_mode,
            transport=args.image_transport, max_workers=args.max_workers,
        )
    else:
        with tempfile.TemporaryDirectory(prefix="rs_caption_tmp_") as td:
            tmp_root = Path(td)
            print(f"[INFO] Using temp dir: {tmp_root}")
            seen, emitted = _process_stream_concurrent(
                stream, total_files, tmp_root, reference_docs, 
                api_url=args.url, model=args.model, max_side=args.max_side, out_jsonl=out_path,
                temperature=args.temperature, top_p=args.top_p, max_tokens=args.max_tokens,
                max_bytes=args.max_bytes, max_pixels=args.max_pixels, detail=args.detail,
                system_prompt=system_prompt, user_prompt_template=user_prompt_template,
                verbose=args.verbose, max_retries=args.max_retries,
                hsi_panels=args.hsi_panels, hsi_triplets=args.hsi_triplets,
                band_indexing=args.band_indexing, clip_percent=args.clip_percent,
                max_panels=args.max_panels, processed=processed, json_mode=args.json_mode,
                transport=args.image_transport, max_workers=args.max_workers,
            )

    print(f"[DONE] Scanned {seen}, wrote {emitted} -> {out_path}")

if __name__ == "__main__":
    main()


